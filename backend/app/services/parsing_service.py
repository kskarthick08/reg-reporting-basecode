"""Parsing helpers for uploaded artifacts and synthetic data loads."""

import json
import math
import re
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from pypdf import PdfReader
from sqlalchemy import text

from app.db import engine

_MODEL_HEADER_LABELS = {
    "table",
    "number of attributes",
    "notes",
    "column name",
    "data type",
    "pk/fk",
    "nullable (y/n)",
    "nullable",
    "psd008 field ref",
    "psd field ref",
    "description",
    "source system",
    "(general)",
}


def chunk_text(text: str, size: int = 1200, overlap: int = 150) -> list[str]:
    """Split long extracted text into overlapping chunks for later context retrieval."""
    text = (text or "").strip()
    if not text:
        return []
    chunks = []
    step = max(1, size - overlap)
    for i in range(0, len(text), step):
        chunk = text[i : i + size].strip()
        if chunk:
            chunks.append(chunk)
    return chunks


def extract_text_from_file(path: Path) -> str:
    """Extract a text preview from supported document formats without mutating the source file."""
    ext = path.suffix.lower()
    if ext == ".txt":
        return read_text_with_fallback(path)
    if ext == ".docx":
        doc = Document(path)
        lines: list[str] = []
        for p in doc.paragraphs:
            txt = (p.text or "").strip()
            if txt:
                lines.append(txt)
        for table in doc.tables:
            for row in table.rows:
                cells = [re.sub(r"\s+", " ", (c.text or "").strip()) for c in row.cells]
                row_txt = " | ".join([c for c in cells if c])
                if row_txt:
                    lines.append(row_txt)
        return "\n".join(lines)
    if ext == ".pdf":
        reader = PdfReader(str(path))
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    if ext in {".csv", ".json", ".xsd", ".xml"}:
        return read_text_with_fallback(path)
    if ext in {".xlsx", ".xls"}:
        try:
            xls = pd.ExcelFile(path)
            blocks: list[str] = []
            for name in xls.sheet_names[:12]:
                df = pd.read_excel(path, sheet_name=name).head(150)
                cols = [str(c).strip() for c in df.columns if str(c).strip()]
                blocks.append(f"[Sheet: {name}]")
                if cols:
                    blocks.append("Columns: " + ", ".join(cols[:120]))
                if not df.empty:
                    preview = df.fillna("").astype(str).to_csv(index=False)
                    blocks.append(preview[:12000])
            return "\n".join(blocks)
        except Exception:
            return ""
    return ""


def strip_nul_text(value: str | None) -> str | None:
    """Handle strip nul text within the service layer."""
    if value is None:
        return None
    return str(value).replace("\x00", "")


def strip_nul_recursive(value: Any):
    """Handle strip nul recursive within the service layer."""
    if isinstance(value, dict):
        return {k: strip_nul_recursive(v) for k, v in value.items()}
    if isinstance(value, list):
        return [strip_nul_recursive(v) for v in value]
    if value is None:
        return None
    try:
        if pd.isna(value):
            return None
    except Exception:
        pass
    if isinstance(value, float) and math.isnan(value):
        return None
    if isinstance(value, str):
        return value.replace("\x00", "")
    return value


def read_text_with_fallback(path: Path) -> str:
    """Read text with fallback within the service layer."""
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return path.read_text(encoding=enc, errors="ignore")
        except Exception:
            continue
    return ""


def _is_noise_model_candidate(value: str) -> bool:
    """Return whether the extracted model value is noisy metadata."""
    txt = re.sub(r"\s+", " ", str(value or "").strip())
    if not txt:
        return True
    lower = txt.lower()
    if lower in _MODEL_HEADER_LABELS:
        return True
    if lower.startswith("example added field:"):
        return True
    if re.fullmatch(r"\d+", txt):
        return True
    if ":" in txt:
        left, right = [part.strip() for part in txt.split(":", 1)]
        if not left or not right:
            return True
        if re.fullmatch(r"\d+", right):
            return True
        if right.lower() in _MODEL_HEADER_LABELS:
            return True
        return False
    if re.fullmatch(r"(bridge|dim|fact|stg|tbl|map)_[a-z0-9_]+", lower):
        return True
    return False


def extract_model_catalog(path: Path) -> dict:
    """Build a normalized catalog of candidate model fields from spreadsheets, JSON, CSV, or SQL."""
    ext = path.suffix.lower()
    if ext in {".xlsx", ".xls"}:
        xls = pd.ExcelFile(path)
        sheets = {}
        raw_headers: list[str] = []
        candidates: list[str] = []
        seen: set[str] = set()

        def _add_candidate(v: str):
            """Add candidate within the service layer."""
            txt = re.sub(r"\s+", " ", str(v or "").strip())
            if not txt:
                return
            if _is_noise_model_candidate(txt):
                return
            key = txt.lower()
            if key in seen:
                return
            seen.add(key)
            candidates.append(txt)

        hint_tokens = ("field", "column", "attribute", "element", "item", "name", "code")
        for name in xls.sheet_names[:20]:
            df = pd.read_excel(path, sheet_name=name)
            cols = [str(c).strip() for c in df.columns if str(c).strip()]
            sheet_table = safe_sql_name(name, fallback="table")
            sheets[name] = cols
            for c in cols:
                raw_headers.append(c)
                _add_candidate(c)

            lower_col_map = {str(c).strip().lower(): c for c in cols}
            model_cols = [lower_col_map[k] for k in lower_col_map if any(h in k for h in hint_tokens)]
            for mc in model_cols[:8]:
                try:
                    vals = df[mc].dropna().astype(str).head(1500).tolist()
                except Exception:
                    vals = []
                for val in vals:
                    raw = re.sub(r"\s+", " ", val.strip())
                    if not raw or raw.lower() in {"nan", "none", "-"}:
                        continue
                    if len(raw) <= 120:
                        _add_candidate(raw)

            table_col = next((lower_col_map[k] for k in lower_col_map if "table" in k), None)
            column_col = next((lower_col_map[k] for k in lower_col_map if any(x in k for x in ("column", "field", "attribute", "element"))), None)
            if table_col and column_col:
                pair_df = df[[table_col, column_col]].dropna().astype(str).head(2000)
                for _, prow in pair_df.iterrows():
                    t = re.sub(r"\s+", " ", prow.get(table_col, "").strip())
                    c = re.sub(r"\s+", " ", prow.get(column_col, "").strip())
                    if t and c:
                        _add_candidate(f"{t}:{c}")
            elif column_col:
                # Common logical-model format: one table per sheet, column names in a dedicated column.
                col_vals = df[column_col].dropna().astype(str).head(3000).tolist()
                for raw_val in col_vals:
                    c = re.sub(r"\s+", " ", raw_val.strip())
                    if not c or c.lower() in {"nan", "none", "-"}:
                        continue
                    if len(c) <= 120 and re.fullmatch(r"[A-Za-z_][A-Za-z0-9_]{1,120}", c):
                        _add_candidate(f"{sheet_table}:{c}")

            try:
                flat = pd.read_excel(path, sheet_name=name, header=None, dtype=str)
            except Exception:
                flat = None
            if flat is not None and not flat.empty:
                cells = flat.fillna("").astype(str).values.flatten().tolist()
                for cell in cells[:8000]:
                    token = re.sub(r"\s+", " ", cell.strip())
                    if not token:
                        continue
                    if re.fullmatch(r"[A-Za-z_][A-Za-z0-9_]{2,}", token):
                        _add_candidate(token)
                    elif re.fullmatch(r"[A-Za-z_][A-Za-z0-9_]{1,80}:[A-Za-z_][A-Za-z0-9_]{1,80}", token):
                        _add_candidate(token)
        combined = candidates
        deduped = []
        seen2 = set()
        for v in combined:
            k = str(v).strip().lower()
            if not k or k in seen2:
                continue
            if _is_noise_model_candidate(v):
                continue
            seen2.add(k)
            deduped.append(str(v).strip())
        return {"sheets": sheets, "fields": deduped[:5000], "targets": deduped[:5000], "headers": raw_headers[:5000]}
    if ext == ".csv":
        df = pd.read_csv(path, nrows=1000)
        cols = [str(c).strip() for c in df.columns if str(c).strip() and not _is_noise_model_candidate(str(c).strip())]
        return {"fields": cols, "targets": cols, "headers": cols}
    if ext == ".json":
        txt = read_text_with_fallback(path)
        payload = json.loads(txt or "{}")
        if isinstance(payload, dict) and isinstance(payload.get("tables"), list):
            # Converted logical model schema format.
            fields: list[str] = []
            targets: list[str] = []
            headers: list[str] = []
            seen: set[str] = set()

            def _add(v: str):
                """Add add within the service layer."""
                vv = re.sub(r"\s+", " ", str(v or "").strip())
                if not vv:
                    return
                k = vv.lower()
                if k in seen:
                    return
                seen.add(k)
                fields.append(vv)
                targets.append(vv)

            for t in payload.get("tables") or []:
                if not isinstance(t, dict):
                    continue
                table = str(t.get("table_name") or t.get("name") or "").strip()
                if not table:
                    continue
                headers.append(table)
                cols = t.get("columns") or []
                for col in cols:
                    if isinstance(col, dict):
                        cname = str(col.get("name") or col.get("source_name") or "").strip()
                    else:
                        cname = str(col or "").strip()
                    if not cname:
                        continue
                    if _is_noise_model_candidate(cname):
                        continue
                    _add(cname)
                    _add(f"{table}:{cname}")
            return {"fields": fields[:5000], "targets": targets[:5000], "headers": headers[:5000]}
        if isinstance(payload, list) and payload and isinstance(payload[0], dict):
            cols = list(payload[0].keys())
            return {"fields": cols, "targets": cols, "headers": cols}
        if isinstance(payload, dict):
            cols = list(payload.keys())
            return {"fields": cols, "targets": cols, "headers": cols}
    if ext == ".sql":
        txt = read_text_with_fallback(path)
        fields: list[str] = []
        targets: list[str] = []
        headers: list[str] = []
        seen: set[str] = set()
        for m in re.finditer(r"create\s+table\s+(?:if\s+not\s+exists\s+)?([A-Za-z_][A-Za-z0-9_]*)\s*\((.*?)\);", txt or "", flags=re.IGNORECASE | re.DOTALL):
            table = str(m.group(1) or "").strip()
            body = str(m.group(2) or "")
            if not table:
                continue
            headers.append(table)
            for ln in body.splitlines():
                part = ln.strip().rstrip(",")
                if not part:
                    continue
                if re.match(r"^(primary|foreign|constraint|unique|check)\b", part, flags=re.IGNORECASE):
                    continue
                cm = re.match(r"^\"?([A-Za-z_][A-Za-z0-9_]*)\"?\b", part)
                if not cm:
                    continue
                col = cm.group(1)
                for vv in (col, f"{table}:{col}"):
                    k = vv.lower()
                    if k in seen:
                        continue
                    seen.add(k)
                    fields.append(vv)
                    targets.append(vv)
        if fields:
            return {"fields": fields[:5000], "targets": targets[:5000], "headers": headers[:5000]}
    txt = read_text_with_fallback(path)
    cols = re.findall(r"[A-Za-z_][A-Za-z0-9_]{2,}", txt or "")
    uniq = []
    seen = set()
    for c in cols:
        lc = c.lower()
        if lc in seen:
            continue
        seen.add(lc)
        uniq.append(c)
    return {"fields": uniq[:5000], "targets": uniq[:5000], "headers": uniq[:5000]}


def extract_data_rows(path: Path) -> list[dict]:
    """Read a bounded sample of tabular source data so downstream prompts stay deterministic."""
    ext = path.suffix.lower()
    if ext in {".xlsx", ".xls"}:
        df = pd.read_excel(path)
        return strip_nul_recursive(df.head(1000).to_dict(orient="records"))
    if ext == ".csv":
        df = pd.read_csv(path)
        return strip_nul_recursive(df.head(1000).to_dict(orient="records"))
    if ext == ".json":
        txt = read_text_with_fallback(path)
        payload = json.loads(txt or "[]")
        if isinstance(payload, dict):
            payload = [payload]
        if isinstance(payload, list):
            return strip_nul_recursive(payload[:1000])
    return []


def extract_mapping_contract(path: Path) -> dict[str, Any] | None:
    """Load an admin-managed JSON mapping contract when the uploaded file matches the expected format."""
    ext = path.suffix.lower()
    if ext != ".json":
        return None
    txt = read_text_with_fallback(path)
    payload = json.loads(txt or "{}")
    if not isinstance(payload, dict):
        return None
    return strip_nul_recursive(payload)


def safe_sql_name(name: str, fallback: str = "table") -> str:
    """Handle safe SQL name within the service layer."""
    n = re.sub(r"[^A-Za-z0-9_]+", "_", str(name or "").strip().lower()).strip("_")
    if not n:
        n = fallback
    if n[0].isdigit():
        n = f"t_{n}"
    return n[:60]


def load_synthetic_folder_to_db(folder_path: Path, table_prefix: str = "synthetic_") -> list[dict[str, Any]]:
    """Materialize supported flat files into PostgreSQL tables for local test data scenarios."""
    folder_path = folder_path.resolve()
    if not folder_path.exists() or not folder_path.is_dir():
        raise FileNotFoundError(f"synthetic_folder_not_found: {folder_path}")

    files = [p for p in folder_path.iterdir() if p.is_file()]
    if not files:
        raise FileNotFoundError(f"no_files_in_synthetic_folder: {folder_path}")

    loaded = []
    with engine.begin() as conn:
        for f in files:
            ext = f.suffix.lower()
            if ext not in {".csv", ".xlsx", ".xls", ".json"}:
                continue
            if ext == ".csv":
                df = pd.read_csv(f)
            elif ext in {".xlsx", ".xls"}:
                df = pd.read_excel(f)
            else:
                txt = read_text_with_fallback(f)
                payload = json.loads(txt or "[]")
                if isinstance(payload, dict):
                    payload = [payload]
                df = pd.DataFrame(payload)

            if df.empty:
                continue

            tname = safe_sql_name(f"{table_prefix}{f.stem}", fallback="synthetic_data")
            cols = [safe_sql_name(c, fallback="col") for c in df.columns]
            df.columns = cols
            conn.execute(text(f'DROP TABLE IF EXISTS "{tname}"'))
            col_defs = ", ".join([f'"{c}" text NULL' for c in cols])
            conn.execute(text(f'CREATE TABLE "{tname}" ({col_defs})'))

            records = df.fillna("").astype(str).to_dict(orient="records")
            if records:
                placeholders = ", ".join([f":{c}" for c in cols])
                col_list = ", ".join([f'"{c}"' for c in cols])
                ins = text(f'INSERT INTO "{tname}" ({col_list}) VALUES ({placeholders})')
                conn.execute(ins, records)

            loaded.append({"table": tname, "rows": len(records), "columns": cols})

    if not loaded:
        raise ValueError("no_supported_synthetic_files_found")
    return loaded
